import os
import json
import requests
import tempfile
import base64
import datetime
from firebase_functions import https_fn
from firebase_functions import options
from google.cloud import storage
from docx import Document
import firebase_functions
from google.auth import impersonated_credentials
from google.auth.transport.requests import Request
import google.auth
import re

# Define project constants first
PROJECT_ID = 'explo-website-tools'
SERVICE_ACCOUNT_EMAIL = '142797649545-compute@developer.gserviceaccount.com'

# Get OpenAI API key from Secret Manager
def get_openai_api_key():
    """Load OpenAI API key from various sources"""
    try:
        # First try environment variables (Firebase auto-creates these from config)
        api_key = (os.environ.get('FUNCTIONS_CONFIG_openai_key') or 
                  os.environ.get('OPENAI_KEY') or 
                  os.environ.get('FUNCTIONS_CONFIG_OPENAI_KEY'))
        
        if api_key:
            print(f"[ADMIN-LOG] ✓ API key loaded from environment variables")
            return api_key
        
        # Try Secret Manager
        try:
            from google.cloud import secretmanager
            
            # Initialize Secret Manager client
            secret_client = secretmanager.SecretManagerServiceClient()
            
            # Access the secret
            secret_name = f"projects/{PROJECT_ID}/secrets/openai-api-key/versions/latest"
            response = secret_client.access_secret_version(request={"name": secret_name})
            api_key = response.payload.data.decode("UTF-8")
            
            print(f"[ADMIN-LOG] ✓ API key loaded from Secret Manager")
            return api_key
            
        except Exception as secret_error:
            print(f"[ADMIN-LOG] Secret Manager failed: {str(secret_error)}")
        
        # Try Firebase Functions config (the working method from before)
        try:
            import firebase_functions
            config = firebase_functions.config()
            api_key = config.get('openai', {}).get('key')
            if api_key:
                print(f"[ADMIN-LOG] ✓ API key loaded from Firebase config")
                return api_key
        except Exception as config_error:
            print(f"[ADMIN-LOG] Firebase config failed: {str(config_error)}")
        
        # Fallback: try environment variable for Firebase config (legacy)
        try:
            api_key = os.environ.get('OPENAI_API_KEY')
            if api_key:
                print(f"[ADMIN-LOG] ✓ API key loaded from OPENAI_API_KEY env var")
                return api_key
        except Exception as env_error:
            print(f"[ADMIN-LOG] Environment variable fallback failed: {str(env_error)}")
        
        print(f"[ADMIN-LOG] ✗ No API key found in any source")
        return None
        
    except Exception as e:
        print(f"[ADMIN-LOG] Error loading OpenAI API key: {e}")
        return None

# Load the API key
OPENAI_API_KEY = get_openai_api_key()

print(f"OpenAI API Key status: {'✓ Loaded' if OPENAI_API_KEY else '✗ Missing'}")
if OPENAI_API_KEY:
    print(f"API Key length: {len(OPENAI_API_KEY)} characters")
    print(f"API Key starts with: {OPENAI_API_KEY[:10]}...")

def generate_secure_signed_url(bucket_name: str, blob_name: str, expiration_hours: int = 2) -> str:
    """
    Generate a secure signed URL using IAM Credentials API.
    
    Args:
        bucket_name: The storage bucket name
        blob_name: The blob/file name
        expiration_hours: Hours until URL expires (default: 2)
    
    Returns:
        str: Signed URL for secure file access
        
    Raises:
        Exception: If signing fails
    """
    try:
        print(f"[ADMIN-LOG] Generating signed URL for: {blob_name}")
        
        # Get default credentials (from Cloud Functions environment)
        source_credentials, _ = google.auth.default()
        
        # Create impersonated credentials for signing
        target_credentials = impersonated_credentials.Credentials(
            source_credentials=source_credentials,
            target_principal=SERVICE_ACCOUNT_EMAIL,
            target_scopes=['https://www.googleapis.com/auth/cloud-platform'],
            delegates=[]
        )
        
        # Initialize storage client with impersonated credentials
        client = storage.Client(credentials=target_credentials)
        bucket = client.bucket(bucket_name)
        blob = bucket.blob(blob_name)
        
        # Generate signed URL with expiration
        expiration = datetime.datetime.utcnow() + datetime.timedelta(hours=expiration_hours)
        signed_url = blob.generate_signed_url(
            expiration=expiration,
            method='GET',
            credentials=target_credentials
        )
        
        print(f"[ADMIN-LOG] ✓ Signed URL generated successfully, expires in {expiration_hours}h")
        return signed_url
        
    except Exception as e:
        print(f"[ADMIN-LOG] ✗ Failed to generate signed URL: {str(e)}")
        raise Exception(f"URL signing failed: {str(e)}")

@https_fn.on_request(cors=options.CorsOptions(cors_origins="*", cors_methods=["GET", "POST"]))
def chat_simple(req: https_fn.Request) -> https_fn.Response:
    """New simplified chat function - direct OpenAI API calls"""
    
    # Handle CORS preflight
    if req.method == 'OPTIONS':
        headers = {
            'Access-Control-Allow-Origin': '*',
            'Access-Control-Allow-Methods': 'POST, OPTIONS',
            'Access-Control-Allow-Headers': 'Content-Type',
            'Access-Control-Max-Age': '3600'
        }
        return https_fn.Response('', headers=headers, status=204)
    
    headers = {
        'Access-Control-Allow-Origin': '*', 
        'Content-Type': 'application/json'
    }
    
    try:
        # Parse request
        request_json = req.get_json(silent=True)
        if not request_json:
            return https_fn.Response(json.dumps({'error': 'Invalid JSON'}), headers=headers, status=400)
            
        conversation = request_json.get('conversation', [])
        if not conversation:
            return https_fn.Response(json.dumps({'error': 'No conversation provided'}), headers=headers, status=400)
        
        # Check if OpenAI API key is available
        if not OPENAI_API_KEY:
            error_msg = 'OpenAI API key not configured. Please set up Firebase config with: firebase functions:config:set openai.key="your-api-key"'
            print(f"ERROR: {error_msg}")
            return https_fn.Response(json.dumps({'error': error_msg}), headers=headers, status=500)
        
        # Prepare conversation for OpenAI API format (much simpler!)
        openai_messages = []
        for msg in conversation[-20:]:  # Last 20 messages for context
            if msg['role'] in ['system', 'user', 'assistant']:
                openai_messages.append({
                    'role': msg['role'],
                    'content': msg['content']
                })
        
        # Call OpenAI API directly
        try:
            openai_response = requests.post(
                'https://api.openai.com/v1/chat/completions',
                headers={
                    'Content-Type': 'application/json',
                    'Authorization': f'Bearer {OPENAI_API_KEY}'
                },
                json={
                    'model': 'gpt-4.1-mini',
                    'messages': openai_messages,
                    'max_tokens': 2048,
                    'temperature': 0.7
                },
                timeout=30  # Add timeout
            )
        except requests.exceptions.RequestException as e:
            error_msg = f"Failed to connect to OpenAI API: {str(e)}"
            print(f"ERROR: {error_msg}")
            return https_fn.Response(json.dumps({'error': error_msg}), headers=headers, status=500)
        
        if openai_response.status_code == 200:
            result = openai_response.json()
            assistant_message = result['choices'][0]['message']['content']
            
            # Extract token usage information
            token_usage = result.get('usage', {})
            
            print(f"[ADMIN-LOG] ✓ Chat response generated. Tokens: {token_usage.get('prompt_tokens', 'N/A')} prompt + {token_usage.get('completion_tokens', 'N/A')} completion = {token_usage.get('total_tokens', 'N/A')} total")
            
            return https_fn.Response(json.dumps({
                'response': assistant_message,
                'tokenUsage': {
                    'promptTokens': token_usage.get('prompt_tokens', 0),
                    'completionTokens': token_usage.get('completion_tokens', 0),
                    'totalTokens': token_usage.get('total_tokens', 0)
                }
            }), headers=headers, status=200)
        else:
            error_msg = f"OpenAI API error: {openai_response.status_code}"
            try:
                error_detail = openai_response.json()
                error_msg += f" - {error_detail.get('error', {}).get('message', '')}"
            except:
                pass
            return https_fn.Response(json.dumps({'error': error_msg}), headers=headers, status=500)
            
    except Exception as e:
        return https_fn.Response(json.dumps({'error': f'Server error: {str(e)}'}), headers=headers, status=500)

@https_fn.on_request(cors=options.CorsOptions(cors_origins="*", cors_methods=["GET", "POST"]))
def export_simple(req: https_fn.Request) -> https_fn.Response:
    """New simplified PRD export function using OpenAI"""
    
    # Handle CORS preflight
    if req.method == 'OPTIONS':
        headers = {
            'Access-Control-Allow-Origin': '*',
            'Access-Control-Allow-Methods': 'POST, OPTIONS',
            'Access-Control-Allow-Headers': 'Content-Type',
            'Access-Control-Max-Age': '3600'
        }
        return https_fn.Response('', headers=headers, status=204)
    
    headers = {
        'Access-Control-Allow-Origin': '*',
        'Content-Type': 'application/json'
    }
    
    try:
        # Parse request
        request_json = req.get_json(silent=True)
        if not request_json:
            return https_fn.Response(json.dumps({'error': 'Invalid JSON'}), headers=headers, status=400)
            
        conversation = request_json.get('conversation', [])
        if not conversation:
            return https_fn.Response(json.dumps({'error': 'No conversation provided'}), headers=headers, status=400)
        
        # Check if OpenAI API key is available
        if not OPENAI_API_KEY:
            error_msg = 'OpenAI API key not configured. Please set up Firebase config with: firebase functions:config:set openai.key="your-api-key"'
            print(f"ERROR: {error_msg}")
            return https_fn.Response(json.dumps({'error': error_msg}), headers=headers, status=500)
        
        # Load accumulated PRD data from storage
        accumulated_prd = {}
        accumulated_summary = ""
        try:
            client = storage.Client()
            bucket = client.bucket(f'{PROJECT_ID}.firebasestorage.app')
            
            # Load PRD data
            prd_blob = bucket.blob('prd_data/current_prd.json')
            if prd_blob.exists():
                existing_json = prd_blob.download_as_text()
                existing_data = json.loads(existing_json)
                accumulated_prd = existing_data.get('sections', {})
                print(f"[ADMIN-LOG] Loaded accumulated PRD with {len(accumulated_prd)} sections for export")
            
            # Load conversation summary
            summary_blob = bucket.blob('prd_data/conversation_summary.txt')
            if summary_blob.exists():
                accumulated_summary = summary_blob.download_as_text()
                print(f"[ADMIN-LOG] Loaded accumulated summary ({len(accumulated_summary)} chars) for export")
            
            if not accumulated_prd and not accumulated_summary:
                print(f"[ADMIN-LOG] No accumulated data found, using conversation only")
                
        except Exception as e:
            print(f"[ADMIN-LOG] Could not load accumulated data: {str(e)}")
        
        # Get recent conversation (messages since last optimization)
        recent_conversation_text = "\n".join([
            f"{msg['role'].upper()}: {msg['content']}" 
            for msg in conversation 
            if msg['role'] != 'system'
        ])
        
        # Create enhanced PRD generation prompt using all accumulated data + recent context
        if accumulated_prd or accumulated_summary:
            context_parts = []
            if accumulated_prd:
                context_parts.append(f"ACCUMULATED PRD SECTIONS:\n{json.dumps(accumulated_prd, indent=2)}")
            if accumulated_summary:
                context_parts.append(f"CONVERSATION HISTORY SUMMARY:\n{accumulated_summary}")
            if recent_conversation_text.strip():
                context_parts.append(f"RECENT CONVERSATION (since last optimization):\n{recent_conversation_text}")
            else:
                context_parts.append("RECENT CONVERSATION: No new messages since last optimization")
            prd_context = f"""{chr(10).join(context_parts)}\n\nINSTRUCTIONS: Generate a comprehensive PRD using all the accumulated data above as the foundation. The PRD sections contain specific structured information, while the conversation summary provides broader context. Incorporate any relevant insights from recent conversation. Prioritize structured PRD data but enhance with conversational context.\n\nIMPORTANT: Do NOT include any reviewer signature tables or signature checklists. In the Problem Statement section, do NOT use boxes, borders, or tables—just use plain text/paragraphs for the content."""
        else:
            prd_context = f"""CONVERSATION DATA:\n{recent_conversation_text}"""
        
        prd_prompt = prd_context + """

Generate a comprehensive Product Requirements Document (PRD) in markdown format using these exact sections in this order:

Use these exact sections in this order:

# Product Requirements Document

## Executive Summary
Write a brief overview of what we're building and why.

## Problem Statement  
What problem are we solving? What pain points exist? *Pain, opportunity, urgency.*

## Goals & Objectives
What are we trying to achieve? What success looks like?
| Type         | Description                                      |
|--------------|--------------------------------------------------|
| **Business** | e.g. "+20 % activated teams within 14 days"       |
| **User**     | e.g. "Get weekly usage insights without leaving app" |
| **Non-Goals**| Explicitly out of scope                          |


## Target Users & Personas
Who will use this product? What are their characteristics?
"As a **[persona]**, I want to **[do X]** so I can **[achieve Y]**."  
• Happy Path Flow → …  
• Edge Cases / Role-Based Flows → …  
• *If visibility is key, prompt for: "Should this experience include analytics, dashboards, or reporting?"*


## User Stories
Key user journeys and use cases in "As a [user], I want [goal] so that [benefit]" format.
1. First impression / entry point  
2. Core flow (include wireframes if possible)  
3. Empty / error / success states  
4. Permissions / access roles  
5. Mobile / a11y considerations


## Features & Requirements
Detailed description of features and functionality we need to build.

## Technical Considerations
High-level technical requirements, constraints, or architecture notes.

## Success Metrics
How will we measure if this is successful? What KPIs matter?
| Metric            | Target | Time Window       |
|-------------------|--------|-------------------|
| **Primary**       |        |                   |
| Secondary         |        |                   |
| CX / Qual Signals |        |                   |

## 🛠️ Technical Considerations
*APIs, data model changes, privacy, scalability, migrations, 3rd-party tools.*  
→ *If querying or embedding analytics is needed, flag potential for live datasets + visualization embedding via Explo.*


## Timeline & Milestones
| Phase   | Scope / Deliverable            | Owner | ETA |
|---------|--------------------------------|-------|-----|
| MVP     |                                |       |     |
| Beta    |                                |       |     |
| GA      |                                |       |     |
| Future  |                                |       |     |

## Out of Scope
What we are NOT building in this version.

## FAQs

Optional: Include an FAQ when helpful to answer high level questions so it is easier for people to grasp the point of the project without getting lost in the details of product definition. 

Impact Checklist

* Permissions  
* Reporting  
* Pricing  
* API  
* Global
"""

        # Get PRD from OpenAI
        try:
            openai_response = requests.post(
                'https://api.openai.com/v1/chat/completions',
                headers={
                    'Content-Type': 'application/json',
                    'Authorization': f'Bearer {OPENAI_API_KEY}'
                },
                json={
                    'model': 'gpt-4.1',
                    'messages': [{'role': 'user', 'content': prd_prompt}],
                    'max_tokens': 4096,
                    'temperature': 0.45
                },
                timeout=60  # Increased timeout to 60 seconds
            )
        except requests.exceptions.RequestException as e:
            error_msg = f"Failed to connect to OpenAI API: {str(e)}"
            print(f"ERROR: {error_msg}")
            return https_fn.Response(json.dumps({'error': error_msg}), headers=headers, status=500)
        
        if openai_response.status_code != 200:
            return https_fn.Response(json.dumps({'error': 'Failed to generate PRD'}), headers=headers, status=500)
            
        result = openai_response.json()
        if 'choices' not in result or len(result['choices']) == 0:
            return https_fn.Response(json.dumps({'error': 'No PRD generated'}), headers=headers, status=500)
            
        prd_markdown = result['choices'][0]['message']['content']
        # Remove leading triple backticks (```markdown or ```) if present
        prd_markdown = prd_markdown.strip()
        if prd_markdown.startswith('```markdown'):
            prd_markdown = prd_markdown[9:].strip()
            if prd_markdown.startswith('\n'):
                prd_markdown = prd_markdown[1:]
            if prd_markdown.endswith('```'):
                prd_markdown = prd_markdown[:-3].strip()
        elif prd_markdown.startswith('```'):
            prd_markdown = prd_markdown[3:].strip()
            if prd_markdown.startswith('\n'):
                prd_markdown = prd_markdown[1:]
            if prd_markdown.endswith('```'):
                prd_markdown = prd_markdown[:-3].strip()
        
        # Convert markdown to Word document
        doc = Document()
        doc.add_heading('Product Requirements Document', 0)
        doc.add_paragraph('Generated by Explo Chat-PRD')
        # Removed page break to avoid empty first page

        def add_markdown_paragraph(md_line):
            """Add a paragraph with bold/italic markdown formatting (all occurrences)."""
            p = doc.add_paragraph()
            pos = 0
            # Find all **bold** and *italic* (not overlapping)
            for match in re.finditer(r'(\*\*[^*]+\*\*|\*[^*]+\*)', md_line):
                start, end = match.span()
                if start > pos:
                    p.add_run(md_line[pos:start])
                text = match.group(0)
                if text.startswith('**') and text.endswith('**'):
                    run = p.add_run(text[2:-2])
                    run.bold = True
                elif text.startswith('*') and text.endswith('*'):
                    run = p.add_run(text[1:-1])
                    run.italic = True
                pos = end
            if pos < len(md_line):
                p.add_run(md_line[pos:])
            return p

        def add_markdown_to_cell(cell, text):
            """Apply markdown bold/italic to a table cell."""
            cell.text = ''  # Clear default
            p = cell.paragraphs[0]
            pos = 0
            for match in re.finditer(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text):
                start, end = match.span()
                if start > pos:
                    p.add_run(text[pos:start])
                mtext = match.group(0)
                if mtext.startswith('**') and mtext.endswith('**'):
                    run = p.add_run(mtext[2:-2])
                    run.bold = True
                elif mtext.startswith('*') and mtext.endswith('*'):
                    run = p.add_run(mtext[1:-1])
                    run.italic = True
                pos = end
            if pos < len(text):
                p.add_run(text[pos:])

        # Table parsing helpers
        def is_table_divider(line):
            return re.match(r'^\|?\s*-+\s*\|', line)
        def is_table_row(line):
            return line.strip().startswith('|') and line.strip().endswith('|')
        
        lines = prd_markdown.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif (line.startswith('- ') or line.startswith('* ')):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif is_table_row(line):
                # Parse markdown table
                header_cells = [cell.strip() for cell in line.strip('|').split('|')]
                i += 1
                # Skip divider
                while i < len(lines) and is_table_divider(lines[i]):
                    i += 1
                # Collect rows
                table_rows = []
                while i < len(lines) and is_table_row(lines[i]):
                    row_cells = [cell.strip() for cell in lines[i].strip('|').split('|')]
                    table_rows.append(row_cells)
                    i += 1
                # Add table to docx
                table = doc.add_table(rows=1, cols=len(header_cells))
                table.style = 'Table Grid'
                for j, cell in enumerate(header_cells):
                    add_markdown_to_cell(table.cell(0, j), cell)
                for row in table_rows:
                    row_cells = table.add_row().cells
                    for j, cell in enumerate(row):
                        add_markdown_to_cell(row_cells[j], cell)
                continue  # already incremented i
            else:
                add_markdown_paragraph(line)
            i += 1
        
        # Save to temporary file
        temp_filename = f"Explo_PRD_{os.urandom(8).hex()}.docx"
        temp_path = os.path.join(tempfile.gettempdir(), temp_filename)
        doc.save(temp_path)
        
        # Upload to Firebase Storage with comprehensive error handling
        try:
            print(f"[ADMIN-LOG] Starting file upload: {temp_filename}")
            
            # Standard storage client for upload (doesn't need signing permissions)
            client = storage.Client()
            bucket_name = f'{PROJECT_ID}.firebasestorage.app'
            bucket = client.bucket(bucket_name)
            blob_path = f'exports/{temp_filename}'
            blob = bucket.blob(blob_path)
            
            # Upload file with metadata for admin tracking
            blob.metadata = {
                'generated_at': datetime.datetime.utcnow().isoformat(),
                'conversation_length': str(len(conversation)),
                'source': 'chat-prd-export'
            }
            blob.upload_from_filename(temp_path)
            
            print(f"[ADMIN-LOG] ✓ File uploaded successfully: {blob_path}")
            print(f"[ADMIN-LOG] File size: {blob.size} bytes")
            
            # Generate secure signed URL
            download_url = generate_secure_signed_url(
                bucket_name=bucket_name,
                blob_name=blob_path,
                expiration_hours=4  # 4-hour expiration for user convenience
            )
            
            print(f"[ADMIN-LOG] ✓ Export completed successfully for user")
            
        except Exception as storage_error:
            print(f"[ADMIN-LOG] ✗ Storage operation failed: {str(storage_error)}")
            # Clean up temp file before returning error
            if os.path.exists(temp_path):
                os.remove(temp_path)
            return https_fn.Response(
                json.dumps({'error': f'File storage failed: {str(storage_error)}'}), 
                headers=headers, 
                status=500
            )
        
        # Clean up temporary file
        try:
            os.remove(temp_path)
            print(f"[ADMIN-LOG] ✓ Temporary file cleaned up")
        except Exception as cleanup_error:
            print(f"[ADMIN-LOG] ⚠ Temp file cleanup failed: {str(cleanup_error)}")
        
        # Return secure download URL
        return https_fn.Response(json.dumps({
            'downloadURL': download_url,
            'fileName': temp_filename,
            'expiresIn': '4 hours'
        }), headers=headers, status=200)
        
    except Exception as e:
        return https_fn.Response(json.dumps({'error': f'Export error: {str(e)}'}), headers=headers, status=500)

@https_fn.on_request(cors=options.CorsOptions(cors_origins="*", cors_methods=["GET", "POST"]))
def optimize_conversation(req: https_fn.Request) -> https_fn.Response:
    """Token optimization: Extract PRD data, update storage, and summarize conversation"""
    
    # Handle CORS preflight
    if req.method == 'OPTIONS':
        headers = {
            'Access-Control-Allow-Origin': '*',
            'Access-Control-Allow-Methods': 'POST, OPTIONS',
            'Access-Control-Allow-Headers': 'Content-Type',
            'Access-Control-Max-Age': '3600'
        }
        return https_fn.Response('', headers=headers, status=204)
    
    headers = {
        'Access-Control-Allow-Origin': '*',
        'Content-Type': 'application/json'
    }
    
    try:
        print(f"[ADMIN-LOG] Starting conversation optimization")
        
        # Parse request
        request_json = req.get_json(silent=True)
        if not request_json:
            return https_fn.Response(json.dumps({'error': 'Invalid JSON'}), headers=headers, status=400)
            
        conversation = request_json.get('conversation', [])
        total_tokens = request_json.get('totalTokens', 0)
        
        if not conversation or len(conversation) <= 1:
            return https_fn.Response(json.dumps({'error': 'No conversation to optimize'}), headers=headers, status=400)
        
        # Check if OpenAI API key is available
        if not OPENAI_API_KEY:
            error_msg = 'OpenAI API key not configured for optimization'
            print(f"ERROR: {error_msg}")
            return https_fn.Response(json.dumps({'error': error_msg}), headers=headers, status=500)
        
        # Extract conversation text (excluding system messages for analysis)
        conversation_text = "\n".join([
            f"{msg['role'].upper()}: {msg['content']}" 
            for msg in conversation 
            if msg['role'] != 'system'
        ])
        
        # Estimate original conversation tokens
        original_tokens = estimate_conversation_tokens(conversation)
        
        print(f"[ADMIN-LOG] Optimizing conversation with {len(conversation)} messages, ~{original_tokens} tokens, user reported {total_tokens} tokens")
        
        # Step 1: Extract PRD-relevant information
        prd_data = extract_prd_information(conversation_text)
        
        # Step 2: Update PRD storage file
        update_prd_storage(prd_data, total_tokens)
        
        # Step 3: Create conversation summary
        conversation_summary = summarize_conversation(conversation_text)
        
        # Step 4: Create optimized conversation with system prompt + summary
        system_prompt = conversation[0] if conversation[0]['role'] == 'system' else {
            'role': 'system',
            'content': 'You are an expert Product Manager AI assistant designed to help users build Product Requirements Documents (PRDs).'
        }
        
        optimized_conversation = [
            system_prompt,
            {
                'role': 'assistant',
                'content': f"""**Previous Conversation Summary:**

{conversation_summary}

**Current PRD Status:** I have extracted and stored information from our previous discussion. Let's continue building your PRD!"""
            }
        ]
        
        # Estimate optimized conversation tokens
        optimized_tokens = estimate_conversation_tokens(optimized_conversation)
        
        print(f"[ADMIN-LOG] ✓ Optimization complete: {len(conversation)} → {len(optimized_conversation)} messages, ~{original_tokens} → ~{optimized_tokens} tokens")
        
        return https_fn.Response(json.dumps({
            'optimizedConversation': optimized_conversation,
            'originalMessages': len(conversation),
            'optimizedMessages': len(optimized_conversation),
            'originalTokens': max(original_tokens, total_tokens),  # Use the higher estimate
            'optimizedTokens': optimized_tokens,
            'tokenSavings': max(original_tokens, total_tokens) - optimized_tokens,
            'prdDataExtracted': len(prd_data),
            'summary': conversation_summary[:200] + "..." if len(conversation_summary) > 200 else conversation_summary
        }), headers=headers, status=200)
        
    except Exception as e:
        print(f"[ADMIN-LOG] ✗ Optimization failed: {str(e)}")
        return https_fn.Response(json.dumps({'error': f'Optimization error: {str(e)}'}), headers=headers, status=500)

def extract_prd_information(conversation_text: str) -> dict:
    """Extract PRD-relevant information from conversation using AI, intelligently merging with existing PRD"""
    
    # Load existing PRD data first
    existing_prd = {}
    try:
        client = storage.Client()
        bucket = client.bucket(f'{PROJECT_ID}.firebasestorage.app')
        blob = bucket.blob('prd_data/current_prd.json')
        
        if blob.exists():
            existing_json = blob.download_as_text()
            existing_data = json.loads(existing_json)
            existing_prd = existing_data.get('sections', {})
            print(f"[ADMIN-LOG] Loaded existing PRD with {len(existing_prd)} sections for intelligent merge")
        else:
            print(f"[ADMIN-LOG] No existing PRD found, starting fresh")
    except Exception as e:
        print(f"[ADMIN-LOG] Could not load existing PRD: {str(e)}")
    
    # Create smart extraction prompt that includes existing PRD context
    prd_extraction_prompt = f"""You are updating an existing PRD with new information from a conversation. 

EXISTING PRD SECTIONS:
{json.dumps(existing_prd, indent=2) if existing_prd else "No existing PRD data"}

NEW CONVERSATION TO ANALYZE:
{conversation_text}

INSTRUCTIONS:
1. Review the EXISTING PRD sections above
2. Analyze the NEW CONVERSATION for relevant information
3. For each PRD section below, intelligently merge new info with existing info:
   - If new info conflicts with existing, prioritize NEW information
   - If new info adds to existing, combine them intelligently
   - If no new info for a section, return the existing content unchanged
   - If neither exists, return null

PRD SECTIONS TO UPDATE:
- executiveSummary: Brief overview of what we're building
- problemStatement: Problems being solved, pain points
- goals: Objectives and success criteria  
- targetUsers: User personas, segments, characteristics
- userStories: User journeys, use cases, "As a user" stories
- features: Specific features and functionality
- technicalConsiderations: Tech requirements, constraints, architecture
- successMetrics: KPIs, measurement criteria
- timeline: Milestones, deadlines, phases
- outOfScope: What we're NOT building

Return ONLY valid JSON with the COMPLETE updated sections (not just changes):
{{"sectionName": "complete updated content or existing content or null"}}

JSON:"""

    try:
        response = requests.post(
            'https://api.openai.com/v1/chat/completions',
            headers={
                'Content-Type': 'application/json',
                'Authorization': f'Bearer {OPENAI_API_KEY}'
            },
            json={
                'model': 'gpt-4.1-mini',
                'messages': [{'role': 'user', 'content': prd_extraction_prompt}],
                'max_tokens': 2048,
                'temperature': 0.15
            },
            timeout=30
        )
        
        if response.status_code == 200:
            result = response.json()
            extracted_text = result['choices'][0]['message']['content']
            
            # Clean and parse JSON
            extracted_text = extracted_text.strip()
            if extracted_text.startswith('```json'):
                extracted_text = extracted_text[7:-3]
            elif extracted_text.startswith('```'):
                extracted_text = extracted_text[3:-3]
                
            return json.loads(extracted_text)
        else:
            print(f"[ADMIN-LOG] PRD extraction failed: {response.status_code}")
            return {}
            
    except Exception as e:
        print(f"[ADMIN-LOG] PRD extraction error: {str(e)}")
        return {}

def summarize_conversation(conversation_text: str) -> str:
    """Create a cumulative summary building on previous summary + new conversation"""
    
    # Load existing summary from storage
    existing_summary = ""
    try:
        client = storage.Client()
        bucket = client.bucket(f'{PROJECT_ID}.firebasestorage.app')
        blob = bucket.blob('prd_data/conversation_summary.txt')
        
        if blob.exists():
            existing_summary = blob.download_as_text()
            print(f"[ADMIN-LOG] Loaded existing summary ({len(existing_summary)} chars) for intelligent merge")
        else:
            print(f"[ADMIN-LOG] No existing summary found, creating first summary")
    except Exception as e:
        print(f"[ADMIN-LOG] Could not load existing summary: {str(e)}")
    
    # Create intelligent summary merging prompt
    if existing_summary.strip():
        summary_prompt = f"""You are updating a cumulative conversation summary with new information.

EXISTING SUMMARY:
{existing_summary}

NEW CONVERSATION SINCE LAST SUMMARY:
{conversation_text}

INSTRUCTIONS:
1. Review the EXISTING SUMMARY above  
2. Analyze the NEW CONVERSATION for additional relevant information
3. Create an UPDATED SUMMARY that intelligently merges both:
   - If new info conflicts with existing, prioritize NEW information
   - If new info adds to existing, incorporate it seamlessly  
   - If new info repeats existing, don't duplicate
   - Maintain key decisions, requirements, and context from both
   - Keep the summary concise but comprehensive (max 1000 words)

Focus on:
- Key decisions and requirements discussed (old + new)
- User needs and problems identified (old + new)
- Features and functionality mentioned (old + new)  
- Technical constraints or preferences (old + new)
- Business goals and success criteria (old + new)

Keep bullets short; preserve explicit numbers, dates, decisions.
UPDATED SUMMARY:"""
    else:
        summary_prompt = f"""Summarize this conversation into key points for PRD context. Focus on:

- Key decisions and requirements discussed
- User needs and problems identified  
- Features and functionality mentioned
- Technical constraints or preferences
- Business goals and success criteria

Keep it concise but comprehensive (max 1000 words).

Keep bullets short; preserve explicit numbers, dates, decisions.

CONVERSATION:
{conversation_text}

SUMMARY:"""

    try:
        response = requests.post(
            'https://api.openai.com/v1/chat/completions',
            headers={
                'Content-Type': 'application/json',
                'Authorization': f'Bearer {OPENAI_API_KEY}'
            },
            json={
                'model': 'gpt-4.1-mini',
                'messages': [{'role': 'user', 'content': summary_prompt}],
                'max_tokens': 1024,
                'temperature': 0.15
            },
            timeout=30
        )
        
        if response.status_code == 200:
            result = response.json()
            updated_summary = result['choices'][0]['message']['content']
            
            # Save updated summary back to storage
            try:
                client = storage.Client()
                bucket = client.bucket(f'{PROJECT_ID}.firebasestorage.app')
                blob = bucket.blob('prd_data/conversation_summary.txt')
                blob.upload_from_string(updated_summary, content_type='text/plain')
                print(f"[ADMIN-LOG] ✓ Updated summary saved ({len(updated_summary)} chars)")
            except Exception as e:
                print(f"[ADMIN-LOG] ✗ Failed to save updated summary: {str(e)}")
            
            return updated_summary
        else:
            print(f"[ADMIN-LOG] Summary generation failed: {response.status_code}")
            return existing_summary if existing_summary else "Previous conversation covered PRD planning and requirements."
            
    except Exception as e:
        print(f"[ADMIN-LOG] Summary error: {str(e)}")
        return existing_summary if existing_summary else "Previous conversation covered PRD planning and requirements."

def update_prd_storage(prd_data: dict, total_tokens: int):
    """Update PRD storage file with new information"""
    
    try:
        # Use Firebase Storage for PRD data persistence
        client = storage.Client()
        bucket = client.bucket(f'{PROJECT_ID}.firebasestorage.app')
        blob = bucket.blob('prd_data/current_prd.json')
        
        # Try to load existing data
        existing_data = {}
        try:
            if blob.exists():
                existing_json = blob.download_as_text()
                existing_data = json.loads(existing_json)
                print(f"[ADMIN-LOG] Loaded existing PRD data: {len(existing_data.get('sections', {}))} sections")
        except Exception as e:
            print(f"[ADMIN-LOG] No existing PRD data found: {str(e)}")
        
        # Merge new data with existing
        updated_data = {
            'lastUpdated': datetime.datetime.utcnow().isoformat(),
            'totalTokens': total_tokens,
            'version': existing_data.get('version', 0) + 1,
            'sections': existing_data.get('sections', {})
        }
        
        # Update each section with new information
        for section, new_info in prd_data.items():
            if new_info and new_info.strip() and new_info.lower() != 'null':
                # Replace the section entirely with AI-merged content (no appending)
                updated_data['sections'][section] = new_info
                print(f"[ADMIN-LOG] Updated section '{section}' with merged content")
        
        # Save updated data
        blob.upload_from_string(
            json.dumps(updated_data, indent=2),
            content_type='application/json'
        )
        
        print(f"[ADMIN-LOG] ✓ PRD data updated: version {updated_data['version']}, {len(updated_data['sections'])} sections")
        
    except Exception as e:
        print(f"[ADMIN-LOG] ✗ PRD storage update failed: {str(e)}")

def estimate_conversation_tokens(conversation: list) -> int:
    """Estimate the total number of tokens in a conversation"""
    
    token_count = 0
    for msg in conversation:
        token_count += len(msg['content'].split())
    return token_count 