# Chat-PRD Streaming API - Version 1.0.1
import os
import json
import asyncio
import time
import tempfile
import logging
import re
import requests
import datetime
from typing import AsyncGenerator, List, Dict, Any
from datetime import datetime

from fastapi import FastAPI, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
import openai
from google.cloud import firestore, secretmanager, storage
from docx import Document
from collections import defaultdict
from dotenv import load_dotenv

# Load environment variables from .env file for local development
load_dotenv()

# Project constants  
PROJECT_ID = os.environ.get('GOOGLE_CLOUD_PROJECT', 'explo-website-tools')
SERVICE_ACCOUNT_EMAIL = '142797649545-compute@developer.gserviceaccount.com'

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(title="Chat-PRD Streaming API", version="1.0.0")

# CORS configuration
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Configure appropriately for production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize Google Cloud services
try:
    db = firestore.Client()
    secret_client = secretmanager.SecretManagerServiceClient()
    storage_client = storage.Client()
    logger.info("Google Cloud services initialized successfully")
except Exception as e:
    logger.error(f"Failed to initialize Google Cloud services: {e}")
    db = None
    secret_client = None
    storage_client = None

# Initialize OpenAI
openai.api_key = os.getenv("OPENAI_API_KEY")

# Rate limiting
class RateLimiter:
    def __init__(self, max_requests: int = 100, window_seconds: int = 60):
        self.max_requests = max_requests
        self.window_seconds = window_seconds
        self.requests = defaultdict(list)
    
    def is_allowed(self, client_ip: str) -> bool:
        now = time.time()
        client_requests = self.requests[client_ip]
        
        # Remove old requests
        client_requests[:] = [req_time for req_time in client_requests 
                             if now - req_time < self.window_seconds]
        
        if len(client_requests) >= self.max_requests:
            return False
        
        client_requests.append(now)
        return True

rate_limiter = RateLimiter()

# Pydantic models
class ChatMessage(BaseModel):
    role: str
    content: str

class ChatRequest(BaseModel):
    conversation: List[ChatMessage]
    user_id: str = None

class ChatResponse(BaseModel):
    response: str
    tokenUsage: Dict[str, int] = None

class ExportRequest(BaseModel):
    conversation: List[ChatMessage]

class OptimizeRequest(BaseModel):
    conversation: List[ChatMessage]
    totalTokens: int = 0

# Utility functions
async def get_openai_key():
    """Retrieve OpenAI API key from Secret Manager or environment"""
    try:
        # First try environment variable (for local development)
        api_key = os.environ.get('OPENAI_API_KEY')
        if api_key:
            logger.info("Using OpenAI API key from environment variable")
            return api_key
        
        # Try Secret Manager
        if secret_client:
            name = f"projects/{PROJECT_ID}/secrets/openai-api-key/versions/latest"
            response = secret_client.access_secret_version(request={"name": name})
            api_key = response.payload.data.decode("UTF-8")
            logger.info("Using OpenAI API key from Secret Manager")
            return api_key
        
        raise Exception("No OpenAI API key found in environment or Secret Manager")
    except Exception as e:
        logger.error(f"Failed to get OpenAI API key: {e}")
        raise HTTPException(status_code=500, detail=f"API key configuration error: {str(e)}")

def estimate_tokens(text: str) -> int:
    """Rough token estimation"""
    return max(1, len(text) // 4)

# Rate limiting middleware
@app.middleware("http")
async def rate_limit_middleware(request: Request, call_next):
    client_ip = request.client.host if request.client else "unknown"
    
    if request.url.path.startswith("/chat"):
        if not rate_limiter.is_allowed(client_ip):
            raise HTTPException(
                status_code=429, 
                detail="Rate limit exceeded. Please slow down."
            )
    
    response = await call_next(request)
    return response

# Client creation with explicit httpx configuration to avoid proxies parameter
async def create_openai_client():
    """Creates an OpenAI client with explicit httpx configuration."""
    import httpx
    api_key = await get_openai_key()
    
    # Create httpx client explicitly without proxies parameter
    http_client = httpx.AsyncClient(
        timeout=30.0,
        limits=httpx.Limits(max_keepalive_connections=5, max_connections=10)
    )
    
    return openai.AsyncOpenAI(
        api_key=api_key, 
        http_client=http_client
    )

# Streaming function
async def stream_openai_response(messages: List[Dict[str, str]]) -> AsyncGenerator[str, None]:
    """Stream OpenAI response chunks with comprehensive logging"""
    request_id = f"req_{int(time.time())}"
    logger.info(f"Starting stream {request_id}")
    
    try:
        client = await create_openai_client()
        
        start_time = time.time()
        chunk_count = 0
        
        # Create streaming completion
        stream = await client.chat.completions.create(
            model="gpt-4o-mini",
            messages=messages,
            stream=True,
            max_tokens=2048,
            temperature=0.7
        )
        
        full_response = ""
        
        async for chunk in stream:
            if chunk.choices[0].delta.content is not None:
                content = chunk.choices[0].delta.content
                full_response += content
                chunk_count += 1
                
                # Yield individual chunk
                chunk_data = {'type': 'chunk', 'content': content, 'chunk_id': chunk_count, 'request_id': request_id}
                yield f"data: {json.dumps(chunk_data)}\n\n"
        
        duration = time.time() - start_time
        logger.info(f"Stream {request_id} completed in {duration:.2f}s with {chunk_count} chunks")
        
        # Final message with complete response
        complete_data = {
            'type': 'complete', 
            'content': full_response,
            'metadata': {
                'duration': duration,
                'chunk_count': chunk_count,
                'request_id': request_id
            }
        }
        yield f"data: {json.dumps(complete_data)}\n\n"
        
        yield "data: [DONE]\n\n"
        
        # Close client properly with error handling
        try:
            await client.close()
        except Exception as close_error:
            logger.warning(f"Stream {request_id} client close error (non-critical): {close_error}")
        
    except Exception as e:
        logger.error(f"Stream {request_id} failed: {str(e)}")
        error_data = {'type': 'error', 'content': f'Streaming error: {str(e)}', 'request_id': request_id}
        yield f"data: {json.dumps(error_data)}\n\n"
        
        # Ensure client cleanup even on error
        try:
            if 'client' in locals():
                await client.close()
        except Exception as close_error:
            logger.warning(f"Stream {request_id} error cleanup client close (non-critical): {close_error}")

# API endpoints
@app.post("/chat/stream")
async def chat_stream(request: ChatRequest):
    """Streaming chat endpoint using Server-Sent Events"""
    try:
        # Prepare messages for OpenAI (last 20 messages for context)
        messages = []
        for msg in request.conversation[-20:]:
            if msg.role in ['system', 'user', 'assistant']:
                messages.append({
                    'role': msg.role,
                    'content': msg.content
                })
        
        if not messages:
            raise HTTPException(status_code=400, detail="No valid messages in conversation")
        
        return StreamingResponse(
            stream_openai_response(messages),
            media_type="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "Connection": "keep-alive",
                "X-Accel-Buffering": "no",  # Disable Nginx buffering
                "Access-Control-Allow-Origin": "*",
                "Access-Control-Allow-Headers": "*",
            }
        )
    except Exception as e:
        logger.error(f"Streaming endpoint error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/chat", response_model=ChatResponse)
async def chat_fallback(request: ChatRequest):
    """Fallback non-streaming endpoint for compatibility"""
    try:
        client = await create_openai_client()
        
        # Prepare messages
        messages = []
        for msg in request.conversation[-20:]:
            if msg.role in ['system', 'user', 'assistant']:
                messages.append({
                    'role': msg.role,
                    'content': msg.content
                })
        
        if not messages:
            raise HTTPException(status_code=400, detail="No valid messages in conversation")
        
        # Create completion
        response = await client.chat.completions.create(
            model="gpt-4o-mini",
            messages=messages,
            max_tokens=2048,
            temperature=0.7
        )
        
        result = ChatResponse(
            response=response.choices[0].message.content,
            tokenUsage={
                "promptTokens": response.usage.prompt_tokens,
                "completionTokens": response.usage.completion_tokens,
                "totalTokens": response.usage.total_tokens
            }
        )
        
        # Close client properly with error handling
        try:
            await client.close()
        except Exception as close_error:
            logger.warning(f"Chat fallback client close error (non-critical): {close_error}")
        return result
    except Exception as e:
        logger.error(f"Chat fallback error: {e}")
        # Ensure client cleanup even on error
        try:
            if 'client' in locals():
                await client.close()
        except Exception as close_error:
            logger.warning(f"Chat fallback error cleanup client close (non-critical): {close_error}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/export")
async def export_prd(request: ExportRequest):
    """Export PRD functionality (simplified version of existing)"""
    try:
        if len(request.conversation) <= 1:
            raise HTTPException(status_code=400, detail="No conversation to export")
        
        api_key = await get_openai_key()
        client = openai.AsyncOpenAI(api_key=api_key)
        
        # Create PRD generation prompt
        conversation_text = "\n".join([
            f"{msg.role.upper()}: {msg.content}" 
            for msg in request.conversation 
            if msg.role != 'system'
        ])
        
        # Use enhanced prompt structure similar to functions/main.py
        prd_context = f"""CONVERSATION DATA:\n{conversation_text}"""
        
        prd_prompt = prd_context + """

Generate a comprehensive Product Requirements Document (PRD) in markdown format using these exact sections in this order:

Use these exact sections in this order:

# Product Requirements Document

## Executive Summary
Write a brief overview of what we're building and why.

## Problem Statement  
What problem are we solving? What pain points exist? *Pain, opportunity, urgency.*

## Goals & Objectives
What are we trying to achieve? What success looks like?
| Type         | Description                                      |
|--------------|--------------------------------------------------|
| **Business** | e.g. "+20 % activated teams within 14 days"       |
| **User**     | e.g. "Get weekly usage insights without leaving app" |
| **Non-Goals**| Explicitly out of scope                          |


## Target Users & Personas
Who will use this product? What are their characteristics?
"As a **[persona]**, I want to **[do X]** so I can **[achieve Y]**."  
• Happy Path Flow → …  
• Edge Cases / Role-Based Flows → …  
• *If visibility is key, prompt for: "Should this experience include analytics, dashboards, or reporting?"*


## User Stories
Key user journeys and use cases in "As a [user], I want [goal] so that [benefit]" format.
1. First impression / entry point  
2. Core flow (include wireframes if possible)  
3. Empty / error / success states  
4. Permissions / access roles  
5. Mobile / a11y considerations


## Features & Requirements
Detailed description of features and functionality we need to build.

## Technical Considerations
High-level technical requirements, constraints, or architecture notes.

## Success Metrics
How will we measure if this is successful? What KPIs matter?
| Metric            | Target | Time Window       |
|-------------------|--------|-------------------|
| **Primary**       |        |                   |
| Secondary         |        |                   |
| CX / Qual Signals |        |                   |

## 🛠️ Technical Considerations
*APIs, data model changes, privacy, scalability, migrations, 3rd-party tools.*  
→ *If querying or embedding analytics is needed, flag potential for live datasets + visualization embedding via Explo.*


## Timeline & Milestones
| Phase   | Scope / Deliverable            | Owner | ETA |
|---------|--------------------------------|-------|-----|
| MVP     |                                |       |     |
| Beta    |                                |       |     |
| GA      |                                |       |     |
| Future  |                                |       |     |

## Out of Scope
What we are NOT building in this version.

## FAQs

Optional: Include an FAQ when helpful to answer high level questions so it is easier for people to grasp the point of the project without getting lost in the details of product definition. 
"""

        # Generate PRD content
        response = await client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prd_prompt}],
            max_tokens=4096,
            temperature=0.5
        )
        
        prd_content = response.choices[0].message.content
        
        # Close client properly with error handling
        try:
            await client.close()
        except Exception as close_error:
            logger.warning(f"Export error cleanup client close (non-critical): {close_error}")
        
        # Create Word document
        doc = Document()
        # Don't add a separate title page - let the AI content provide the main heading
        
        def add_markdown_paragraph(md_line):
            """Add a paragraph with bold/italic markdown formatting (all occurrences), including for list items."""
            # If this is a bullet point, use List Bullet style
            is_bullet = md_line.startswith('- ') or md_line.startswith('* ')
            text = md_line[2:] if is_bullet else md_line
            p = doc.add_paragraph(style='List Bullet' if is_bullet else None)
            pos = 0
            import re
            for match in re.finditer(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text):
                start, end = match.span()
                if start > pos:
                    p.add_run(text[pos:start])
                mtext = match.group(0)
                if mtext.startswith('**') and mtext.endswith('**'):
                    run = p.add_run(mtext[2:-2])
                    run.bold = True
                elif mtext.startswith('*') and mtext.endswith('*'):
                    run = p.add_run(mtext[1:-1])
                    run.italic = True
                pos = end
            if pos < len(text):
                p.add_run(text[pos:])
            return p

        def add_markdown_to_cell(cell, text):
            """Apply markdown bold/italic to a table cell."""
            cell.text = ''  # Clear default
            p = cell.paragraphs[0]
            pos = 0
            import re
            for match in re.finditer(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text):
                start, end = match.span()
                if start > pos:
                    p.add_run(text[pos:start])
                mtext = match.group(0)
                if mtext.startswith('**') and mtext.endswith('**'):
                    run = p.add_run(mtext[2:-2])
                    run.bold = True
                elif mtext.startswith('*') and mtext.endswith('*'):
                    run = p.add_run(mtext[1:-1])
                    run.italic = True
                pos = end
            if pos < len(text):
                p.add_run(text[pos:])

        # Table parsing helpers
        def is_table_divider(line):
            import re
            return re.match(r'^\|?\s*-+\s*\|', line)
        def is_table_row(line):
            return line.strip().startswith('|') and line.strip().endswith('|')
        
        # Enhanced markdown to Word conversion
        lines = prd_content.split('\n')
        i = 0
        prd_title_added = False  # Track if we've added the generated by text
        
        while i < len(lines):
            line = lines[i].strip()
            # Remove lines that are just '---' (markdown horizontal rules)
            if line == '---':
                i += 1
                continue
            if not line:
                i += 1
                continue
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
                # Add "Generated by Explo Chat-PRD" right after the main title
                if not prd_title_added and 'Product Requirements Document' in line:
                    doc.add_paragraph('Generated by Explo Chat-PRD').alignment = 1  # Center align
                    prd_title_added = True
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif (line.startswith('- ') or line.startswith('* ')):
                add_markdown_paragraph(line)
            elif is_table_row(line):
                # Parse markdown table
                header_cells = [cell.strip() for cell in line.strip('|').split('|')]
                i += 1
                # Skip divider
                while i < len(lines) and is_table_divider(lines[i]):
                    i += 1
                # Collect rows
                table_rows = []
                while i < len(lines) and is_table_row(lines[i]):
                    row_cells = [cell.strip() for cell in lines[i].strip('|').split('|')]
                    table_rows.append(row_cells)
                    i += 1
                # Add table to docx
                table = doc.add_table(rows=1, cols=len(header_cells))
                table.style = 'Table Grid'
                for j, cell in enumerate(header_cells):
                    add_markdown_to_cell(table.cell(0, j), cell)
                for row in table_rows:
                    row_cells = table.add_row().cells
                    for j, cell in enumerate(row):
                        if j < len(row):
                            add_markdown_to_cell(row_cells[j], cell)
                continue  # already incremented i
            else:
                add_markdown_paragraph(line)
            i += 1
        
        # Save to temporary file
        temp_filename = f"PRD_{int(time.time())}.docx"
        temp_path = os.path.join(tempfile.gettempdir(), temp_filename)
        doc.save(temp_path)
        
        # Upload to Google Cloud Storage with public access (no signed URL needed)
        if storage_client:
            try:
                bucket_name = f'{PROJECT_ID}.firebasestorage.app'
                bucket = storage_client.bucket(bucket_name)
                blob_path = f'exports/{temp_filename}'
                blob = bucket.blob(blob_path)
                
                # Upload file
                blob.upload_from_filename(temp_path)
                
                # Make blob publicly readable
                blob.make_public()
                
                # Use public URL instead of signed URL
                download_url = blob.public_url
                
                # Clean up temp file
                os.remove(temp_path)
                
                logger.info(f"File uploaded successfully: {download_url}")
                return {"downloadURL": download_url, "fileName": temp_filename}
                
            except Exception as storage_error:
                logger.error(f"Storage upload failed: {storage_error}")
                
                # Alternative: Return base64 encoded file for small files
                try:
                    import base64
                    with open(temp_path, 'rb') as f:
                        file_data = f.read()
                    
                    # Only if file is small enough (< 1MB)
                    if len(file_data) < 1024 * 1024:
                        encoded_data = base64.b64encode(file_data).decode('utf-8')
                        os.remove(temp_path)  # Clean up temp file
                        
                        return {
                            "downloadData": encoded_data,
                            "fileName": temp_filename,
                            "mimeType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        }
                except Exception as fallback_error:
                    logger.error(f"Fallback encoding failed: {fallback_error}")
                
                # Final fallback: return error with local file reference
                return {"error": "Storage upload failed", "localFile": temp_path}
        else:
            return {"error": "Storage not configured", "localFile": temp_path}
            
    except Exception as e:
        logger.error(f"Export error: {e}")
        # Ensure client cleanup even on error
        try:
            if 'client' in locals():
                await client.close()
        except Exception as close_error:
            logger.warning(f"Export error cleanup client close (non-critical): {close_error}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/optimize")
async def optimize_conversation(request: OptimizeRequest):
    """Token optimization: Extract PRD data, update storage, and summarize conversation"""
    try:
        logger.info("Starting conversation optimization")
        
        if len(request.conversation) <= 1:
            raise HTTPException(status_code=400, detail="No conversation to optimize")
        
        # Convert Pydantic models to dict for processing
        conversation = [{"role": msg.role, "content": msg.content} for msg in request.conversation]
        total_tokens = request.totalTokens
        
        # Extract conversation text (excluding system messages for analysis)
        conversation_text = "\n".join([
            f"{msg['role'].upper()}: {msg['content']}" 
            for msg in conversation 
            if msg['role'] != 'system'
        ])
        
        # Estimate original conversation tokens
        original_tokens = estimate_conversation_tokens(conversation)
        
        logger.info(f"Optimizing conversation with {len(conversation)} messages, ~{original_tokens} tokens, user reported {total_tokens} tokens")
        
        # Step 1: Extract PRD-relevant information
        prd_data = await extract_prd_information(conversation_text)
        
        # Step 2: Update PRD storage file
        await update_prd_storage(prd_data, total_tokens)
        
        # Step 3: Create conversation summary
        conversation_summary = await summarize_conversation(conversation_text)
        
        # Step 4: Create optimized conversation with system prompt + summary
        system_prompt = conversation[0] if conversation[0]['role'] == 'system' else {
            'role': 'system',
            'content': 'You are an expert Product Manager AI assistant designed to help users build Product Requirements Documents (PRDs).'
        }
        
        optimized_conversation = [
            system_prompt,
            {
                'role': 'assistant',
                'content': f"""**Previous Conversation Summary:**

{conversation_summary}

**Current PRD Status:** I have extracted and stored information from our previous discussion. Let's continue building your PRD!"""
            }
        ]
        
        # Estimate optimized conversation tokens
        optimized_tokens = estimate_conversation_tokens(optimized_conversation)
        
        logger.info(f"Optimization complete: {len(conversation)} → {len(optimized_conversation)} messages, ~{original_tokens} → ~{optimized_tokens} tokens")
        
        return {
            'optimizedConversation': [
                {"role": msg["role"], "content": msg["content"]} 
                for msg in optimized_conversation
            ],
            'originalMessages': len(conversation),
            'optimizedMessages': len(optimized_conversation),
            'originalTokens': max(original_tokens, total_tokens),  # Use the higher estimate
            'optimizedTokens': optimized_tokens,
            'tokenSavings': max(original_tokens, total_tokens) - optimized_tokens,
            'prdDataExtracted': len(prd_data),
            'summary': conversation_summary[:200] + "..." if len(conversation_summary) > 200 else conversation_summary
        }
        
    except Exception as e:
        logger.error(f"Optimization failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Optimization error: {str(e)}")

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {"status": "healthy", "timestamp": datetime.now().isoformat()}

@app.get("/warmup")
async def warmup():
    """Warm-up endpoint that initializes basic services for faster first responses"""
    try:
        start_time = time.time()
        logger.info("Starting warmup...")
        
        # Basic service checks that don't require API keys
        services_status = {
            "fastapi": "ready",
            "firestore": "ready" if db else "unavailable", 
            "storage": "ready" if storage_client else "unavailable"
        }
        
        # Test OpenAI API key availability (but don't fail if not available)
        try:
            if os.environ.get('OPENAI_API_KEY'):
                # If API key is in environment, test client creation
                import httpx
                http_client = httpx.AsyncClient(
                    timeout=30.0,
                    limits=httpx.Limits(max_keepalive_connections=5, max_connections=10)
                )
                await http_client.aclose()  # Just test creation and close
                services_status["openai"] = "ready"
                logger.info("OpenAI client initialization test successful")
            else:
                services_status["openai"] = "no_api_key"
        except Exception as openai_error:
            logger.warning(f"OpenAI warmup test failed (non-critical): {openai_error}")
            services_status["openai"] = "error"
        
        warmup_time = time.time() - start_time
        logger.info(f"Warmup completed in {warmup_time:.2f}s")
        
        return {
            "status": "warmed_up", 
            "timestamp": datetime.now().isoformat(),
            "warmup_time": warmup_time,
            "services": services_status
        }
    except Exception as e:
        logger.error(f"Warmup error: {e}")
        return {
            "status": "warmup_failed", 
            "error": str(e),
            "timestamp": datetime.now().isoformat()
        }

@app.get("/")
async def root():
    """Root endpoint"""
    return {"message": "Chat-PRD Streaming API", "version": "1.0.0"}

# Optimization helper functions
def estimate_conversation_tokens(conversation: List[Dict[str, str]]) -> int:
    """Estimate the total number of tokens in a conversation"""
    token_count = 0
    for msg in conversation:
        token_count += len(msg['content'].split())
    return token_count

async def extract_prd_information(conversation_text: str) -> dict:
    """Extract PRD-relevant information from conversation using AI, intelligently merging with existing PRD"""
    
    # Load existing PRD data first
    existing_prd = {}
    try:
        if storage_client:
            bucket_name = f'{PROJECT_ID}.firebasestorage.app'
            bucket = storage_client.bucket(bucket_name)
            blob = bucket.blob('prd_data/current_prd.json')
            
            if blob.exists():
                existing_json = blob.download_as_text()
                existing_data = json.loads(existing_json)
                existing_prd = existing_data.get('sections', {})
                logger.info(f"Loaded existing PRD with {len(existing_prd)} sections for intelligent merge")
            else:
                logger.info("No existing PRD found, starting fresh")
    except Exception as e:
        logger.info(f"Could not load existing PRD: {str(e)}")
    
    # Create smart extraction prompt that includes existing PRD context
    prd_extraction_prompt = f"""You are updating an existing PRD with new information from a conversation. 

EXISTING PRD SECTIONS:
{json.dumps(existing_prd, indent=2) if existing_prd else "No existing PRD data"}

NEW CONVERSATION TO ANALYZE:
{conversation_text}

INSTRUCTIONS:
1. Review the EXISTING PRD sections above
2. Analyze the NEW CONVERSATION for relevant information
3. For each PRD section below, intelligently merge new info with existing info:
   - If new info conflicts with existing, prioritize NEW information
   - If new info adds to existing, combine them intelligently
   - If no new info for a section, return the existing content unchanged
   - If neither exists, return null

PRD SECTIONS TO UPDATE:
- executiveSummary: Brief overview of what we're building
- problemStatement: Problems being solved, pain points
- goals: Objectives and success criteria  
- targetUsers: User personas, segments, characteristics
- userStories: User journeys, use cases, "As a user" stories
- features: Specific features and functionality
- technicalConsiderations: Tech requirements, constraints, architecture
- successMetrics: KPIs, measurement criteria
- timeline: Milestones, deadlines, phases
- outOfScope: What we're NOT building

Return ONLY valid JSON with the COMPLETE updated sections (not just changes):
{{"sectionName": "complete updated content or existing content or null"}}

JSON:"""

    try:
        api_key = await get_openai_key()
        
        response = requests.post(
            'https://api.openai.com/v1/chat/completions',
            headers={
                'Content-Type': 'application/json',
                'Authorization': f'Bearer {api_key}'
            },
            json={
                'model': 'gpt-4o-mini',
                'messages': [{'role': 'user', 'content': prd_extraction_prompt}],
                'max_tokens': 2048,
                'temperature': 0.15
            },
            timeout=30
        )
        
        if response.status_code == 200:
            result = response.json()
            extracted_text = result['choices'][0]['message']['content']
            
            # Clean and parse JSON
            extracted_text = extracted_text.strip()
            if extracted_text.startswith('```json'):
                extracted_text = extracted_text[7:-3]
            elif extracted_text.startswith('```'):
                extracted_text = extracted_text[3:-3]
                
            return json.loads(extracted_text)
        else:
            logger.error(f"PRD extraction failed: {response.status_code}")
            return {}
            
    except Exception as e:
        logger.error(f"PRD extraction error: {str(e)}")
        return {}

async def summarize_conversation(conversation_text: str) -> str:
    """Create a cumulative summary building on previous summary + new conversation"""
    
    # Load existing summary from storage
    existing_summary = ""
    try:
        if storage_client:
            bucket_name = f'{PROJECT_ID}.firebasestorage.app'
            bucket = storage_client.bucket(bucket_name)
            blob = bucket.blob('prd_data/conversation_summary.txt')
            
            if blob.exists():
                existing_summary = blob.download_as_text()
                logger.info(f"Loaded existing summary ({len(existing_summary)} chars) for intelligent merge")
            else:
                logger.info("No existing summary found, creating first summary")
    except Exception as e:
        logger.info(f"Could not load existing summary: {str(e)}")
    
    # Create intelligent summary merging prompt
    if existing_summary.strip():
        summary_prompt = f"""You are updating a cumulative conversation summary with new information.

EXISTING SUMMARY:
{existing_summary}

NEW CONVERSATION SINCE LAST SUMMARY:
{conversation_text}

INSTRUCTIONS:
1. Review the EXISTING SUMMARY above  
2. Analyze the NEW CONVERSATION for additional relevant information
3. Create an UPDATED SUMMARY that intelligently merges both:
   - If new info conflicts with existing, prioritize NEW information
   - If new info adds to existing, incorporate it seamlessly  
   - If new info repeats existing, don't duplicate
   - Maintain key decisions, requirements, and context from both
   - Keep the summary concise but comprehensive (max 1000 words)

Focus on:
- Key decisions and requirements discussed (old + new)
- User needs and problems identified (old + new)
- Features and functionality mentioned (old + new)  
- Technical constraints or preferences (old + new)
- Business goals and success criteria (old + new)

Keep bullets short; preserve explicit numbers, dates, decisions.
UPDATED SUMMARY:"""
    else:
        summary_prompt = f"""Summarize this conversation into key points for PRD context. Focus on:

- Key decisions and requirements discussed
- User needs and problems identified  
- Features and functionality mentioned
- Technical constraints or preferences
- Business goals and success criteria

Keep it concise but comprehensive (max 1000 words).

Keep bullets short; preserve explicit numbers, dates, decisions.

CONVERSATION:
{conversation_text}

SUMMARY:"""

    try:
        api_key = await get_openai_key()
        
        response = requests.post(
            'https://api.openai.com/v1/chat/completions',
            headers={
                'Content-Type': 'application/json',
                'Authorization': f'Bearer {api_key}'
            },
            json={
                'model': 'gpt-4o-mini',
                'messages': [{'role': 'user', 'content': summary_prompt}],
                'max_tokens': 1024,
                'temperature': 0.15
            },
            timeout=30
        )
        
        if response.status_code == 200:
            result = response.json()
            updated_summary = result['choices'][0]['message']['content']
            
            # Save updated summary back to storage
            try:
                if storage_client:
                    bucket_name = f'{PROJECT_ID}.firebasestorage.app'
                    bucket = storage_client.bucket(bucket_name)
                    blob = bucket.blob('prd_data/conversation_summary.txt')
                    blob.upload_from_string(updated_summary, content_type='text/plain')
                    logger.info(f"Updated summary saved ({len(updated_summary)} chars)")
            except Exception as e:
                logger.error(f"Failed to save updated summary: {str(e)}")
            
            return updated_summary
        else:
            logger.error(f"Summary generation failed: {response.status_code}")
            return existing_summary if existing_summary else "Previous conversation covered PRD planning and requirements."
            
    except Exception as e:
        logger.error(f"Summary error: {str(e)}")
        return existing_summary if existing_summary else "Previous conversation covered PRD planning and requirements."

async def update_prd_storage(prd_data: dict, total_tokens: int):
    """Update PRD storage file with new information"""
    
    try:
        if not storage_client:
            logger.warning("Storage client not available")
            return
            
        # Use Firebase Storage for PRD data persistence
        bucket_name = f'{PROJECT_ID}.firebasestorage.app'
        bucket = storage_client.bucket(bucket_name)
        blob = bucket.blob('prd_data/current_prd.json')
        
        # Try to load existing data
        existing_data = {}
        try:
            if blob.exists():
                existing_json = blob.download_as_text()
                existing_data = json.loads(existing_json)
                logger.info(f"Loaded existing PRD data: {len(existing_data.get('sections', {}))} sections")
        except Exception as e:
            logger.info(f"No existing PRD data found: {str(e)}")
        
        # Merge new data with existing
        updated_data = {
            'lastUpdated': datetime.datetime.utcnow().isoformat(),
            'totalTokens': total_tokens,
            'version': existing_data.get('version', 0) + 1,
            'sections': existing_data.get('sections', {})
        }
        
        # Update each section with new information
        for section, new_info in prd_data.items():
            if new_info and new_info.strip() and new_info.lower() != 'null':
                # Replace the section entirely with AI-merged content (no appending)
                updated_data['sections'][section] = new_info
                logger.info(f"Updated section '{section}' with merged content")
        
        # Save updated data
        blob.upload_from_string(
            json.dumps(updated_data, indent=2),
            content_type='application/json'
        )
        
        logger.info(f"PRD data updated: version {updated_data['version']}, {len(updated_data['sections'])} sections")
        
    except Exception as e:
        logger.error(f"PRD storage update failed: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 8080))
    uvicorn.run(app, host="0.0.0.0", port=port) 